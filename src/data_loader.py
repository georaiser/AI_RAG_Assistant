# data_loader.py
"""
Document loading with improved error handling and debugging.
"""

import os
import logging
from pathlib import Path
from typing import List, Dict, Any, Tuple, Optional
import pdfplumber
import pandas as pd
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangchainDocument
from config import Config

logger = logging.getLogger(__name__)

class DocumentLoader:
    """Document loader with improved error handling and debugging."""
    
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=Config.CHUNK_SIZE,
            chunk_overlap=Config.CHUNK_OVERLAP,
            separators=["\n\n", "\n", " ", ""]
        )
        self.loaded_files = []
        self.file_stats = {}
        self.processing_errors = []
    
    def load_documents(self) -> List[LangchainDocument]:
        """Load documents based on configuration."""
        if Config.DEBUG_MODE:
            logger.info("Starting document loading")
        
        if Config.PROCESS_MULTIPLE_DOCUMENTS:
            documents = self._load_multiple_documents()
        else:
            documents = self._load_single_document()
        
        logger.info(f"Loaded {len(self.loaded_files)} files, {len(documents)} chunks")
        
        if Config.VERBOSE_MODE and documents:
            # Log sample of loaded content
            for i, doc in enumerate(documents[:3]):
                source = doc.metadata.get('source', 'Unknown')
                page = doc.metadata.get('page', 'N/A')
                content_preview = doc.page_content[:100].replace('\n', ' ')
                logger.info(f"Sample {i+1}: {source} (Page: {page}) - {content_preview}...")
        
        return documents
    
    def get_file_info(self) -> Dict[str, Any]:
        """Get file information with actual loaded files list."""
        total_chunks = sum(stats.get("chunks", 0) for stats in self.file_stats.values())
        
        return {
            "loaded_files": self.loaded_files.copy(),
            "file_stats": self.file_stats,
            "total_files": len(self.loaded_files),
            "total_chunks": total_chunks,
            "processing_errors": self.processing_errors
        }
    
    def _load_single_document(self) -> List[LangchainDocument]:
        """Load single document."""
        if not Config.SINGLE_DOCUMENT_PATH.exists():
            error_msg = f"Document not found: {Config.SINGLE_DOCUMENT_PATH}"
            logger.error(error_msg)
            self.processing_errors.append(error_msg)
            return []
        
        return self._process_file(Config.SINGLE_DOCUMENT_PATH)
    
    def _load_multiple_documents(self) -> List[LangchainDocument]:
        """Load multiple documents with better progress tracking."""
        if not Config.DOCUMENTS_DIR.exists():
            error_msg = f"Documents directory not found: {Config.DOCUMENTS_DIR}"
            logger.error(error_msg)
            self.processing_errors.append(error_msg)
            return []
        
        all_documents = []
        supported_files = [
            f for f in Config.DOCUMENTS_DIR.rglob("*") 
            if self._is_supported_file(f)
        ]
        
        if not supported_files:
            error_msg = f"No supported files found in {Config.DOCUMENTS_DIR}"
            logger.warning(error_msg)
            self.processing_errors.append(error_msg)
            return []
        
        if Config.DEBUG_MODE:
            logger.info(f"Found {len(supported_files)} supported files")
        
        for i, file_path in enumerate(supported_files):
            try:
                documents = self._process_file(file_path)
                all_documents.extend(documents)
                
                if Config.DEBUG_MODE:
                    logger.info(f"[{i+1}/{len(supported_files)}] Processed {file_path.name}: {len(documents)} chunks")
                    
            except Exception as e:
                error_msg = f"Failed to process {file_path.name}: {str(e)}"
                logger.error(error_msg)
                self.processing_errors.append(error_msg)
        
        return all_documents
    
    def _is_supported_file(self, file_path: Path) -> bool:
        """Check if file is supported with better validation."""
        if not file_path.is_file():
            return False
            
        if file_path.suffix.lower() not in Config.SUPPORTED_FORMATS:
            return False
            
        if file_path.name.startswith('.'):
            return False
            
        try:
            size = file_path.stat().st_size
            if size == 0:
                if Config.VERBOSE_MODE:
                    logger.warning(f"Skipping empty file: {file_path.name}")
                return False
            if size > 100 * 1024 * 1024:  # 100MB limit
                logger.warning(f"Skipping large file (>100MB): {file_path.name}")
                return False
        except OSError:
            return False
            
        return True
    
    def _process_file(self, file_path: Path) -> List[LangchainDocument]:
        """Process a single file with improved error handling."""
        try:
            documents = []
            
            # Route to appropriate processor
            if file_path.suffix.lower() == '.pdf':
                documents = self._process_pdf(file_path)
            elif file_path.suffix.lower() in ['.docx', '.dotx']:
                documents = self._process_docx(file_path)
            else:
                documents = self._process_text_file(file_path)
            
            # Update tracking
            self.loaded_files.append(file_path.name)
            self.file_stats[file_path.name] = {
                "chunks": len(documents),
                "size_mb": round(file_path.stat().st_size / (1024 * 1024), 3),
                "file_type": file_path.suffix.lower()
            }
            
            return documents
            
        except Exception as e:
            error_msg = f"Error processing {file_path.name}: {str(e)}"
            logger.error(error_msg)
            self.processing_errors.append(error_msg)
            return []
    
    def _add_citation_header(self, content: str, filename: str, page: Optional[int] = None) -> str:
        """Add citation header - always include page when available."""
        if page is not None:
            header = f"[{filename}, Page: {page}]"
        else:
            header = f"[{filename}]"
        
        return f"{header}\n{content}"
    
    def _process_pdf(self, file_path: Path) -> List[LangchainDocument]:
        """Process PDF with improved error handling."""
        documents = []
        
        try:
            with pdfplumber.open(file_path) as pdf:
                total_pages = len(pdf.pages)
                
                if Config.VERBOSE_MODE:
                    logger.info(f"Processing PDF {file_path.name} with {total_pages} pages")
                
                for page_num, page in enumerate(pdf.pages, 1):
                    try:
                        # Extract text content
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            cleaned_text = self._clean_text(page_text)
                            if len(cleaned_text.strip()) < 50:  # Skip very short content
                                continue
                                
                            content_with_header = self._add_citation_header(
                                cleaned_text, file_path.name, page_num
                            )
                            
                            # Split into chunks
                            chunks = self.text_splitter.split_text(content_with_header)
                            
                            for chunk_idx, chunk in enumerate(chunks):
                                # Ensure chunk has proper header
                                if not chunk.startswith(f"[{file_path.name}"):
                                    chunk = self._add_citation_header(chunk, file_path.name, page_num)
                                
                                metadata = {
                                    "source": file_path.name,
                                    "page": page_num,
                                    "chunk_index": chunk_idx,
                                    "content_type": "text",
                                    "file_type": "pdf"
                                }
                                documents.append(LangchainDocument(page_content=chunk, metadata=metadata))
                        
                        # Extract tables with better error handling
                        try:
                            tables = page.extract_tables()
                            for table_idx, table in enumerate(tables):
                                if table and len(table) > 1 and len(table[0]) > 1:
                                    df = pd.DataFrame(table[1:], columns=table[0])
                                    df = df.dropna(how='all').fillna('')
                                    
                                    if not df.empty and len(df) > 0:
                                        table_content = f"Table {table_idx + 1}:\n{df.to_string(index=False)}"
                                        content_with_header = self._add_citation_header(
                                            table_content, file_path.name, page_num
                                        )
                                        
                                        metadata = {
                                            "source": file_path.name,
                                            "page": page_num,
                                            "table_index": table_idx,
                                            "content_type": "table",
                                            "file_type": "pdf"
                                        }
                                        documents.append(LangchainDocument(page_content=content_with_header, metadata=metadata))
                        except Exception as e:
                            if Config.VERBOSE_MODE:
                                logger.debug(f"Error processing table on page {page_num}: {e}")
                                
                    except Exception as e:
                        logger.warning(f"Error processing page {page_num} of {file_path.name}: {e}")
                        continue
        
        except Exception as e:
            logger.error(f"Error reading PDF {file_path.name}: {e}")
            raise
        
        return documents
    
    def _process_docx(self, file_path: Path) -> List[LangchainDocument]:
        """Process DOCX file with improved error handling."""
        documents = []
        
        try:
            doc = Document(file_path)
            
            # Process paragraphs
            paragraphs = []
            for p in doc.paragraphs:
                text = p.text.strip()
                if text and len(text) > 10:  # Skip very short paragraphs
                    paragraphs.append(text)
            
            if paragraphs:
                all_text = "\n".join(paragraphs)
                cleaned_text = self._clean_text(all_text)
                
                if cleaned_text.strip():
                    content_with_header = self._add_citation_header(cleaned_text, file_path.name)
                    chunks = self.text_splitter.split_text(content_with_header)
                    
                    for chunk_idx, chunk in enumerate(chunks):
                        if not chunk.startswith(f"[{file_path.name}"):
                            chunk = self._add_citation_header(chunk, file_path.name)
                        
                        metadata = {
                            "source": file_path.name,
                            "chunk_index": chunk_idx,
                            "content_type": "text",
                            "file_type": "docx"
                        }
                        documents.append(LangchainDocument(page_content=chunk, metadata=metadata))
            
            # Process tables
            for table_idx, table in enumerate(doc.tables):
                try:
                    table_data = []
                    for row in table.rows:
                        row_data = [cell.text.strip() for cell in row.cells]
                        if any(cell for cell in row_data):
                            table_data.append(row_data)
                    
                    if len(table_data) > 1:
                        df = pd.DataFrame(table_data[1:], columns=table_data[0])
                        df = df.dropna(how='all').fillna('')
                        
                        if not df.empty:
                            table_content = f"Table {table_idx + 1}:\n{df.to_string(index=False)}"
                            content_with_header = self._add_citation_header(table_content, file_path.name)
                            
                            metadata = {
                                "source": file_path.name,
                                "table_index": table_idx,
                                "content_type": "table",
                                "file_type": "docx"
                            }
                            documents.append(LangchainDocument(page_content=content_with_header, metadata=metadata))
                except Exception as e:
                    if Config.VERBOSE_MODE:
                        logger.debug(f"Error processing table {table_idx}: {e}")
        
        except Exception as e:
            logger.error(f"Error reading DOCX {file_path.name}: {e}")
            raise
        
        return documents
    
    def _process_text_file(self, file_path: Path) -> List[LangchainDocument]:
        """Process text files with improved encoding detection."""
        documents = []
        
        # Try different encodings
        encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252', 'iso-8859-1']
        content = None
        used_encoding = None
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    content = file.read()
                used_encoding = encoding
                break
            except UnicodeDecodeError:
                continue
        
        if content is None:
            raise ValueError(f"Could not decode {file_path.name} with any encoding")
        
        if Config.VERBOSE_MODE:
            logger.info(f"Decoded {file_path.name} using {used_encoding}")
        
        cleaned_content = self._clean_text(content)
        if cleaned_content.strip() and len(cleaned_content.strip()) > 50:
            content_with_header = self._add_citation_header(cleaned_content, file_path.name)
            chunks = self.text_splitter.split_text(content_with_header)
            
            for chunk_idx, chunk in enumerate(chunks):
                if not chunk.startswith(f"[{file_path.name}"):
                    chunk = self._add_citation_header(chunk, file_path.name)
                
                metadata = {
                    "source": file_path.name,
                    "chunk_index": chunk_idx,
                    "content_type": "text",
                    "file_type": file_path.suffix.lower() or "txt"
                }
                documents.append(LangchainDocument(page_content=chunk, metadata=metadata))
        
        return documents
    
    def _clean_text(self, text: str) -> str:
        """Clean text content with improved handling."""
        if not text:
            return ""
        
        # Split into lines and clean each line
        lines = text.split('\n')
        cleaned_lines = []
        prev_empty = False
        
        for line in lines:
            line = line.strip()
            if line:
                # Remove excessive whitespace within line
                line = ' '.join(line.split())
                cleaned_lines.append(line)
                prev_empty = False
            elif not prev_empty:
                cleaned_lines.append('')
                prev_empty = True
        
        # Remove trailing empty lines
        while cleaned_lines and not cleaned_lines[-1]:
            cleaned_lines.pop()
        
        return '\n'.join(cleaned_lines).strip()


def load_data() -> Tuple[List[LangchainDocument], Dict[str, Any]]:
    """Main function to load documents with improved error handling."""
    try:
        loader = DocumentLoader()
        documents = loader.load_documents()
        file_info = loader.get_file_info()
        
        if Config.DEBUG_MODE:
            logger.info(f"Loading complete: {file_info['total_files']} files, {file_info['total_chunks']} chunks")
            if file_info.get('processing_errors'):
                logger.warning(f"Processing errors: {len(file_info['processing_errors'])}")
        
        return documents, file_info
    
    except Exception as e:
        logger.error(f"Critical error in load_data: {e}")
        return [], {
            "error": str(e), 
            "total_files": 0, 
            "total_chunks": 0,
            "loaded_files": [],
            "processing_errors": [str(e)]
        }