"""
Document loader for single file processing with table extraction.
"""
import logging
from typing import List, Dict, Any
from pathlib import Path
import pdfplumber
import docx
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document
from src.config import config

logger = logging.getLogger(__name__)


def create_chunks(pages_data: List[Dict[str, Any]]) -> List[Document]:
    """Create document chunks with metadata."""
    if not pages_data:
        return []
    
    # Initialize text splitter
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=config.CHUNK_SIZE,
        chunk_overlap=config.CHUNK_OVERLAP,
        separators=config.SEPARATORS,
        length_function=len,
    )
    
    documents = []
    
    for page_data in pages_data:
        # Split text into chunks
        chunks = text_splitter.split_text(page_data["text"])
        
        for i, chunk in enumerate(chunks):
            if not chunk.strip():
                continue
            
            # Create metadata
            metadata = {
                "page": page_data["page"],
                "source": page_data["source"],
                "file_type": page_data.get("file_type", "unknown")
            }
            
            # Create document
            doc = Document(page_content=chunk, metadata=metadata)
            documents.append(doc)
    
    logger.info(f"Created {len(documents)} document chunks")
    return documents


class PDFLoader:
    """PDF document loader with table extraction."""
    
    def load(self, file_path: Path) -> List[Dict[str, Any]]:
        """Load PDF with text and table extraction."""
        logger.info(f"Loading PDF: {file_path}")
        pages_data = []
        
        try:
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    # Extract text
                    text = page.extract_text()
                    if not text or not text.strip():
                        continue
                    
                    # Extract tables and convert to text
                    tables = page.extract_tables()
                    table_text = ""
                    if tables:
                        table_parts = []
                        for i, table in enumerate(tables):
                            if table:
                                # Convert table to simple text format
                                table_rows = []
                                for row in table:
                                    if row:
                                        clean_row = [str(cell).strip() if cell else "" for cell in row]
                                        table_rows.append(" | ".join(clean_row))
                                table_parts.append(f"Table {i+1}:\n" + "\n".join(table_rows))
                        table_text = "\n\n".join(table_parts)
                    
                    # Combine text and tables
                    full_text = text.strip()
                    if table_text:
                        full_text += f"\n\n[TABLES]\n{table_text}"
                    
                    pages_data.append({
                        "text": full_text,
                        "page": page_num,
                        "source": file_path.name,
                        "file_type": "pdf"
                    })
        
        except Exception as e:
            logger.error(f"Error loading PDF {file_path}: {e}")
            raise
        
        logger.info(f"Loaded {len(pages_data)} pages from PDF")
        return pages_data


class TextLoader:
    """Simple text file loader."""
    
    def load(self, file_path: Path) -> List[Dict[str, Any]]:
        """Load plain text files."""
        logger.info(f"Loading text file: {file_path}")
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read().strip()
            
            if not content:
                return []
            
            return [{
                "text": content,
                "page": 1,
                "source": file_path.name,
                "file_type": "text"
            }]
        
        except Exception as e:
            logger.error(f"Error loading text file {file_path}: {e}")
            raise


class DocxLoader:
    """Word document loader with table extraction."""
    
    def load(self, file_path: Path) -> List[Dict[str, Any]]:
        """Load Word documents with table extraction."""
        logger.info(f"Loading Word document: {file_path}")
        
        try:
            doc = docx.Document(str(file_path))
            
            # Extract paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text.strip())
            
            # Extract tables
            table_texts = []
            for i, table in enumerate(doc.tables):
                rows = []
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    rows.append(row_text)
                table_texts.append(f"Table {i+1}:\n" + "\n".join(rows))
            
            # Combine all content
            all_text = "\n\n".join(paragraphs)
            if table_texts:
                all_text += f"\n\n[TABLES]\n" + "\n\n".join(table_texts)
            
            if not all_text.strip():
                return []
            
            return [{
                "text": all_text,
                "page": 1,
                "source": file_path.name,
                "file_type": "docx"
            }]
        
        except Exception as e:
            logger.error(f"Error loading Word document {file_path}: {e}")
            raise


class DocumentLoader:
    """Document loader."""
    
    def __init__(self):
        self.loaders = {
            '.pdf': PDFLoader(),
            '.txt': TextLoader(),
            '.md': TextLoader(),
            '.rst': TextLoader(),
            '.docx': DocxLoader(),
            '.dotx': DocxLoader()
        }
    
    def get_loader(self, file_path: Path):
        """Get appropriate loader for file type."""
        ext = file_path.suffix.lower()
        if ext not in self.loaders:
            supported = list(self.loaders.keys())
            raise ValueError(f"Unsupported file type: {ext}. Supported: {supported}")
        return self.loaders[ext]
    
    def load_and_process_document(self, file_path: Path) -> List[Document]:
        """Load and process single document."""
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Get appropriate loader and load document data
        loader = self.get_loader(file_path)
        pages_data = loader.load(file_path)
        
        # Create chunks
        documents = create_chunks(pages_data)
        
        if not documents:
            logger.warning(f"No documents created from {file_path}")
            return []
        
        logger.info(f"Successfully processed {file_path}: {len(documents)} chunks")
        return documents

    def load_multiple_documents(self, file_paths: List[Path]) -> List[Document]:
        """Load and process multiple documents - ADD THIS."""
        all_documents = []
        for file_path in file_paths:
            try:
                documents = self.load_and_process_document(file_path)
                # Add document metadata
                for doc in documents:
                    doc.metadata['document_id'] = file_path.stem
                    doc.metadata['document_name'] = file_path.name
                all_documents.extend(documents)
            except Exception as e:
                logger.error(f"Failed to load {file_path}: {e}")
                continue
        return all_documents

class DocumentLoaderFactory:
    """Factory for supported extensions."""
    
    @staticmethod
    def supported_extensions() -> List[str]:
        """Return supported file extensions."""
        return ['.pdf', '.txt', '.md', '.rst', '.docx', '.dotx']